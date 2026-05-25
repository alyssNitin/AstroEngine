"""
backend/reports/docx_generator.py
====================================
Generates a styled Microsoft Word (.docx) report from a user's Vedic reading.

Uses python-docx if available; falls back to a plain-text .txt file encoded
as bytes when the library is not installed.

Public API
----------
    generate_docx(session_data, user_profile) -> bytes
        Returns the raw bytes of a .docx (or .txt fallback) file.

    build_report_bytes(session_id, db) -> tuple[bytes, str]
        Convenience wrapper — looks up session + profile from the database
        and returns (file_bytes, filename).
"""
from __future__ import annotations

import io
import re
import textwrap
from datetime import datetime
from typing import Any


# ── Tiny HTML stripper ────────────────────────────────────────────────────────

def _strip_html(text: str) -> str:
    text = re.sub(r"<[^>]+>", "", text or "")
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    text = text.replace("&nbsp;", " ").replace("&#39;", "'").replace("&quot;", '"')
    return text.strip()


def _clean(text: str) -> str:
    return (
        _strip_html(text)
        .replace("‘", "'").replace("’", "'")
        .replace("“", '"').replace("”", '"')
        .replace("—", "--").replace("–", "-")
    )


# ── Plain-text fallback ───────────────────────────────────────────────────────

def _build_plaintext(
    user_name: str,
    birth_info: str,
    refined_analysis: str,
    chat_messages: list[dict],
    dasha_text: str = "",
    overall_theme: str = "",
) -> bytes:
    lines: list[str] = []
    sep = "=" * 72
    thin = "-" * 72

    lines += [sep, "  NARAYAN ASTRO READER — YOUR PERSONALISED VEDIC REPORT", sep]
    lines += [
        f"Prepared for : {_clean(user_name)}",
        f"Birth details: {_clean(birth_info)}",
        f"Generated    : {datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}",
        "",
    ]

    if overall_theme:
        lines += [thin, "OVERALL LIFE THEME", thin]
        for para in _clean(overall_theme).split("\n\n"):
            lines.extend(textwrap.wrap(para.strip(), 72))
        lines.append("")

    lines += [thin, "PART 1 — DEEP VEDIC READING", thin]
    for para in _clean(refined_analysis).split("\n\n"):
        para = para.strip()
        if para:
            lines.extend(textwrap.wrap(para, 72))
            lines.append("")

    if dasha_text:
        lines += [thin, "PART 2 — DASHA PERIODS", thin]
        for para in _clean(dasha_text).split("\n\n"):
            para = para.strip()
            if para:
                lines.extend(textwrap.wrap(para, 72))
                lines.append("")

    q_idx = 1
    for msg in chat_messages:
        role = msg.get("role", "")
        body = _clean(msg.get("content", ""))
        if role == "user" and body and body != "[Prior session restored]":
            if q_idx == 1:
                lines += [thin, "PART 3 — YOUR QUESTIONS & ANSWERS", thin]
            lines.append(f"Q{q_idx}: {body}")
        elif role == "assistant" and body:
            lines.append(f"A{q_idx}: {body}")
            lines.append("")
            q_idx += 1

    lines += [sep, "© NarayanAstroReader — Vedic AI Astrology", sep]
    return "\n".join(lines).encode("utf-8")


# ── python-docx report ────────────────────────────────────────────────────────

def _build_docx(
    user_name: str,
    birth_info: str,
    refined_analysis: str,
    chat_messages: list[dict],
    dasha_text: str = "",
    overall_theme: str = "",
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Colour palette ────────────────────────────────────────────────────────
    DEEP_PURPLE   = RGBColor(0x4A, 0x00, 0x8C)   # headings
    GOLD          = RGBColor(0xC0, 0x96, 0x00)   # sub-headings
    DARK_GREY     = RGBColor(0x22, 0x22, 0x22)   # body text
    MID_GREY      = RGBColor(0x55, 0x55, 0x55)   # secondary text
    Q_BLUE        = RGBColor(0x1A, 0x3F, 0x8F)   # questions
    A_GREEN       = RGBColor(0x14, 0x5A, 0x32)   # answers

    def _set_font(run, size_pt: int, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None):
        run.font.name  = "Palatino Linotype"
        run.font.size  = Pt(size_pt)
        run.bold       = bold
        run.italic     = italic
        if color:
            run.font.color.rgb = color

    def _heading(text: str, level: int, color: RGBColor = DEEP_PURPLE):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        _set_font(run, {1: 20, 2: 15, 3: 12}[level], bold=True, color=color)
        p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _body(text: str, indent: bool = False, color: RGBColor = DARK_GREY):
        p   = doc.add_paragraph()
        run = p.add_run(_clean(text))
        _set_font(run, 11, color=color)
        p.paragraph_format.space_after   = Pt(6)
        p.paragraph_format.line_spacing  = Pt(16)
        if indent:
            p.paragraph_format.left_indent = Cm(1)
        return p

    def _hr():
        """Thin horizontal rule via bottom-border on an empty paragraph."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "9966CC")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # ── Cover ─────────────────────────────────────────────────────────────────
    title_p  = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("NARAYAN ASTRO READER")
    _set_font(run, 26, bold=True, color=DEEP_PURPLE)

    sub_p  = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Your Personalised Vedic Horoscope Report")
    _set_font(run, 14, italic=True, color=GOLD)

    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(
        f"Prepared for: {_clean(user_name)}\n"
        f"{_clean(birth_info)}\n"
        f"Generated: {datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}"
    )
    _set_font(run, 11, color=MID_GREY)

    _hr()

    # ── Overall theme ─────────────────────────────────────────────────────────
    if overall_theme:
        _heading("Overall Life Theme", level=2, color=GOLD)
        for para in _clean(overall_theme).split("\n\n"):
            para = para.strip()
            if para:
                _body(para)
        _hr()

    # ── Deep reading ──────────────────────────────────────────────────────────
    _heading("Part 1 — Deep Vedic Reading", level=1)
    for para in _clean(refined_analysis).split("\n\n"):
        para = para.strip()
        if para:
            # Section headers inside the reading (ALL-CAPS lines)
            if para.isupper() or (len(para) < 60 and para.endswith(":")):
                _heading(para.rstrip(":"), level=3, color=GOLD)
            else:
                _body(para)
    _hr()

    # ── Dasha periods ─────────────────────────────────────────────────────────
    if dasha_text:
        _heading("Part 2 — Dasha (Planetary) Periods", level=1)
        for para in _clean(dasha_text).split("\n\n"):
            para = para.strip()
            if para:
                if para.isupper() or "Dasha" in para or "period" in para.lower():
                    _heading(para, level=3, color=GOLD)
                else:
                    _body(para)
        _hr()

    # ── Chat Q&A ──────────────────────────────────────────────────────────────
    q_pairs: list[tuple[str, str]] = []
    pending_q = ""
    for msg in chat_messages:
        role = msg.get("role", "")
        body = _clean(msg.get("content", ""))
        if role == "user" and body and body != "[Prior session restored]":
            pending_q = body
        elif role == "assistant" and pending_q:
            q_pairs.append((pending_q, body))
            pending_q = ""

    if q_pairs:
        part_num = 3 if dasha_text else 2
        _heading(f"Part {part_num} — Your Questions & Answers", level=1)
        for idx, (q, a) in enumerate(q_pairs, 1):
            qp  = doc.add_paragraph()
            run = qp.add_run(f"Q{idx}: {q}")
            _set_font(run, 11, bold=True, color=Q_BLUE)
            qp.paragraph_format.space_before = Pt(8)
            qp.paragraph_format.space_after  = Pt(2)

            ap  = doc.add_paragraph()
            run = ap.add_run(a)
            _set_font(run, 11, color=A_GREEN)
            ap.paragraph_format.left_indent = Cm(0.5)
            ap.paragraph_format.space_after = Pt(8)
        _hr()

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_p  = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("© NarayanAstroReader — Vedic AI Astrology")
    _set_font(run, 9, italic=True, color=MID_GREY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────

def generate_docx(
    session_data: dict[str, Any],
    user_profile: dict[str, Any] | None = None,
) -> bytes:
    """
    Build a Word report from session_data (as stored in _sessions dict).

    Falls back to plain-text bytes when python-docx is not installed.
    """
    profile      = user_profile or {}
    user_name    = (
        session_data.get("name")
        or profile.get("name")
        or "Valued Member"
    )
    dob  = profile.get("date_of_birth", session_data.get("dob", ""))
    tob  = profile.get("time_of_birth", session_data.get("tob", ""))
    pob  = profile.get("place_of_birth", session_data.get("pob", ""))
    birth_info   = f"{dob}, {tob}, {pob}".strip(", ")

    refined      = session_data.get("refined_analysis", "")
    if not refined:
        refined = profile.get("refined_analysis", "")
    overall      = session_data.get("overall_theme", profile.get("overall_theme", ""))
    messages     = session_data.get("messages", [])
    dasha_text   = session_data.get("dasha_text", "")

    try:
        return _build_docx(user_name, birth_info, refined, messages,
                           dasha_text, overall)
    except ImportError:
        return _build_plaintext(user_name, birth_info, refined, messages,
                                dasha_text, overall)
    except Exception as exc:
        # Last-resort fallback — always return something
        return _build_plaintext(user_name, birth_info, refined, messages,
                                dasha_text, overall)


def build_report_bytes(session_id: str, db, sessions: dict) -> tuple[bytes, str]:
    """
    Look up session + user profile and return (file_bytes, filename).

    Args:
        session_id: active in-memory session key
        db:         Database instance
        sessions:   the _sessions dict from main.py

    Returns:
        (bytes, filename)  — filename ends in .docx or .txt
    """
    session = sessions.get(session_id, {})
    email   = session.get("email", "")
    profile = db.get_profile(email) if email else {}

    data    = generate_docx(session, profile)
    ext     = "docx" if data[:4] == b"PK\x03\x04" else "txt"

    safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_",
                       (profile or {}).get("name", "reading") or "reading")
    filename  = f"vedic_reading_{safe_name}_{datetime.utcnow().strftime('%Y%m%d')}.{ext}"
    return data, filename
